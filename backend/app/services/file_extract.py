"""File extraction service.

Turns user-uploaded attachments (text, source code, PDF, DOCX, images, etc.)
into a normalized text representation that the SentinelGuard threat pipeline
can scan exactly like a regular prompt.

Optional parsers (pypdf, python-docx, pytesseract+Pillow) are imported lazily
so the gateway works out of the box even when those packages are not present;
in that case the attachment is preserved with a graceful "binary not extracted"
note instead of failing the whole request.
"""

from __future__ import annotations

import asyncio
import base64
import binascii
import io
import mimetypes
import os
from dataclasses import dataclass, field

from app.core.config import get_settings
from app.core.logging import get_logger

log = get_logger("file_extract")


MAX_FILES_PER_REQUEST = 8
MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB / file
MAX_TOTAL_BYTES = 30 * 1024 * 1024  # 30 MB / request
MAX_EXTRACTED_CHARS = 60_000  # per file, after extraction


_TEXT_EXTS = {
    ".txt", ".md", ".rst", ".log", ".csv", ".tsv", ".json", ".jsonl",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".env", ".xml", ".html",
    ".htm", ".css", ".scss", ".less",
}

_CODE_EXTS = {
    ".py": "python", ".pyi": "python", ".ipynb": "json",
    ".js": "javascript", ".mjs": "javascript", ".cjs": "javascript",
    ".jsx": "jsx", ".ts": "typescript", ".tsx": "tsx",
    ".java": "java", ".kt": "kotlin", ".scala": "scala",
    ".go": "go", ".rs": "rust",
    ".c": "c", ".h": "c", ".cpp": "cpp", ".cc": "cpp", ".hpp": "cpp",
    ".cs": "csharp", ".swift": "swift",
    ".rb": "ruby", ".php": "php",
    ".sh": "bash", ".bash": "bash", ".zsh": "bash", ".ps1": "powershell",
    ".sql": "sql", ".r": "r", ".dart": "dart", ".lua": "lua",
    ".dockerfile": "dockerfile", ".tf": "hcl", ".hcl": "hcl",
}


@dataclass
class ExtractedAttachment:
    """Normalized representation of an uploaded file."""

    filename: str
    mime_type: str
    size_bytes: int
    kind: str  # text | code | pdf | docx | image | binary
    language: str | None = None
    content_text: str = ""
    truncated: bool = False
    error: str | None = None
    metadata: dict = field(default_factory=dict)
    # Retained for image attachments so post-extraction enrichment (e.g.
    # multimodal vision description) can reuse the bytes without re-decoding
    # the original base64 payload. Excluded from repr / API summaries.
    raw_bytes: bytes | None = field(default=None, repr=False, compare=False)

    def as_prompt_block(self) -> str:
        """Render this attachment as a fenced prompt block.

        The result is appended to the user prompt so every existing scanner
        (regex, PII, secrets, jailbreak embeddings, etc.) sees the file
        contents and can flag prompt-injection or sensitive data hidden
        inside an upload.
        """
        header = f"[Attached file: {self.filename} ({self.mime_type}, {self.size_bytes} bytes, {self.kind})]"
        if self.error:
            return f"{header}\n(content not extracted: {self.error})"
        if self.kind == "image":
            note = self.content_text or "(binary image — no OCR text available)"
            return f"{header}\n{note}"
        if not self.content_text:
            return f"{header}\n(empty)"
        fence_lang = self.language or ""
        body = self.content_text
        if self.truncated:
            body += f"\n\n... [truncated to {MAX_EXTRACTED_CHARS} chars]"
        return f"{header}\n```{fence_lang}\n{body}\n```"

    def to_summary(self) -> dict:
        """Lightweight metadata for API responses / audit logs (no raw content)."""
        return {
            "filename": self.filename,
            "mime_type": self.mime_type,
            "size_bytes": self.size_bytes,
            "kind": self.kind,
            "language": self.language,
            "truncated": self.truncated,
            "extracted_chars": len(self.content_text),
            "error": self.error,
            "has_ocr": bool(self.metadata.get("has_ocr")),
            "has_vision": bool(self.metadata.get("has_vision")),
        }


def _decode_base64(data_b64: str) -> bytes:
    cleaned = data_b64.split(",", 1)[1] if data_b64.startswith("data:") else data_b64
    cleaned = cleaned.strip()
    try:
        return base64.b64decode(cleaned, validate=False)
    except (binascii.Error, ValueError) as e:
        raise ValueError(f"invalid base64 payload: {e}") from e


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text, False
    return text[:MAX_EXTRACTED_CHARS], True


def _classify(filename: str, mime_type: str) -> tuple[str, str | None]:
    """Return (kind, language)."""
    ext = os.path.splitext(filename or "")[1].lower()
    mime = (mime_type or "").lower()

    if ext in _CODE_EXTS:
        return "code", _CODE_EXTS[ext]
    if ext in _TEXT_EXTS or mime.startswith("text/") or mime in {
        "application/json", "application/xml", "application/x-yaml",
        "application/x-toml", "application/javascript",
    }:
        lang = ext.lstrip(".") or "text"
        return "text", lang
    if mime == "application/pdf" or ext == ".pdf":
        return "pdf", None
    if ext == ".docx" or mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return "docx", None
    if mime.startswith("image/") or ext in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}:
        return "image", None
    return "binary", None


def _decode_text_bytes(raw: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_pdf(raw: bytes) -> tuple[str, str | None]:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as e:  # noqa: BLE001
        return "", f"pypdf not installed ({e})"
    try:
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                pages.append(page.extract_text() or "")
            except Exception as pe:  # noqa: BLE001
                pages.append(f"[page {i+1} extraction failed: {pe}]")
        return "\n\n".join(pages).strip(), None
    except Exception as e:  # noqa: BLE001
        return "", f"pdf parse error: {e}"


def _extract_docx(raw: bytes) -> tuple[str, str | None]:
    try:
        import docx  # type: ignore # python-docx
    except Exception as e:  # noqa: BLE001
        return "", f"python-docx not installed ({e})"
    try:
        document = docx.Document(io.BytesIO(raw))
        paragraphs = [p.text for p in document.paragraphs if p.text]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs).strip(), None
    except Exception as e:  # noqa: BLE001
        return "", f"docx parse error: {e}"


def _extract_image_ocr(raw: bytes) -> tuple[str, str | None]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except Exception:  # noqa: BLE001
        return "", None
    try:
        img = Image.open(io.BytesIO(raw))
        text = pytesseract.image_to_string(img) or ""
        text = text.strip()
        if not text:
            return "", None
        return f"[OCR text extracted from image]\n{text}", None
    except Exception as e:  # noqa: BLE001
        return "", f"ocr error: {e}"


def extract_one(
    *,
    filename: str,
    mime_type: str | None,
    data_b64: str,
) -> ExtractedAttachment:
    """Extract a single attachment described by base64 payload."""

    fname = (filename or "upload.bin").strip() or "upload.bin"
    mime = (mime_type or "").strip()
    if not mime:
        guessed, _ = mimetypes.guess_type(fname)
        mime = guessed or "application/octet-stream"

    try:
        raw = _decode_base64(data_b64)
    except ValueError as e:
        return ExtractedAttachment(
            filename=fname, mime_type=mime, size_bytes=0,
            kind="binary", error=str(e),
        )

    size = len(raw)
    if size > MAX_FILE_BYTES:
        return ExtractedAttachment(
            filename=fname, mime_type=mime, size_bytes=size,
            kind="binary",
            error=f"file too large ({size} bytes; limit {MAX_FILE_BYTES})",
        )

    kind, language = _classify(fname, mime)
    att = ExtractedAttachment(
        filename=fname, mime_type=mime, size_bytes=size,
        kind=kind, language=language,
    )

    try:
        if kind in ("text", "code"):
            text, truncated = _truncate(_decode_text_bytes(raw))
            att.content_text = text
            att.truncated = truncated
        elif kind == "pdf":
            text, err = _extract_pdf(raw)
            if err:
                att.error = err
            text, truncated = _truncate(text)
            att.content_text = text
            att.truncated = truncated
        elif kind == "docx":
            text, err = _extract_docx(raw)
            if err:
                att.error = err
            text, truncated = _truncate(text)
            att.content_text = text
            att.truncated = truncated
        elif kind == "image":
            att.raw_bytes = raw
            ocr, err = _extract_image_ocr(raw)
            if err:
                att.metadata["ocr_error"] = err
            text, truncated = _truncate(ocr)
            att.content_text = text
            att.truncated = truncated
            att.metadata["has_ocr"] = bool(text)
        else:
            att.error = "binary file type not supported for extraction"
    except Exception as e:  # noqa: BLE001
        log.warning("extract_failed", filename=fname, mime=mime, error=str(e))
        att.error = f"extraction failed: {e}"

    return att


def extract_many(items: list[dict]) -> tuple[list[ExtractedAttachment], str | None]:
    """Validate count / total size, then extract all attachments.

    Returns (attachments, error). On error, attachments may be partial.
    """
    if not items:
        return [], None
    if len(items) > MAX_FILES_PER_REQUEST:
        return [], f"too many files (max {MAX_FILES_PER_REQUEST})"

    total = 0
    out: list[ExtractedAttachment] = []
    for raw in items:
        att = extract_one(
            filename=str(raw.get("filename") or "upload.bin"),
            mime_type=raw.get("mime_type"),
            data_b64=str(raw.get("data_b64") or ""),
        )
        total += att.size_bytes
        if total > MAX_TOTAL_BYTES:
            att.error = f"total upload size exceeded {MAX_TOTAL_BYTES} bytes"
            att.content_text = ""
            out.append(att)
            return out, att.error
        out.append(att)
    return out, None


def merge_into_prompt(prompt: str, attachments: list[ExtractedAttachment]) -> str:
    """Append rendered attachment blocks to the user's prompt."""
    if not attachments:
        return prompt
    blocks = [a.as_prompt_block() for a in attachments]
    joined = "\n\n".join(blocks)
    if not prompt.strip():
        return joined
    return f"{prompt}\n\n{joined}"


async def enrich_with_vision(attachments: list[ExtractedAttachment]) -> None:
    """Append a multimodal-vision description to each image attachment in-place.

    The description is generated by a vision-capable LLM (configured via
    ``VISION_MODEL``) using a security-focused system prompt that surfaces
    visible text, PII, credentials, URLs / QR codes, NSFW content, etc.
    The resulting text is appended to ``content_text`` so every existing
    scanner (regex, PII, secrets, jailbreak embeddings, toxicity, ...) sees
    the visual content alongside any OCR output.

    No-ops when:
      * ``VISION_DESCRIBE_ENABLED`` is false
      * No image attachments are present
      * No LLM credentials are configured for the chosen model

    Failures per-image are swallowed and logged so the request still
    proceeds with whatever OCR text is already available.
    """
    settings = get_settings()
    if not settings.vision_describe_enabled:
        return

    targets = [a for a in attachments if a.kind == "image" and a.raw_bytes]
    if not targets:
        return

    from app.llm.litellm_client import adescribe_image

    async def _run(att: ExtractedAttachment) -> tuple[ExtractedAttachment, str | None]:
        desc = await adescribe_image(
            raw=att.raw_bytes or b"",
            mime=att.mime_type,
            model=settings.vision_model,
            timeout=settings.vision_timeout,
            max_tokens=settings.vision_max_tokens,
        )
        return att, desc

    results = await asyncio.gather(
        *[_run(a) for a in targets], return_exceptions=True
    )

    for result in results:
        if isinstance(result, BaseException):
            log.warning("vision_enrich_task_failed", error=str(result))
            continue
        att, desc = result
        if not desc:
            continue
        block = f"[Vision description]\n{desc}"
        merged = f"{att.content_text}\n\n{block}".strip() if att.content_text else block
        merged_text, truncated = _truncate(merged)
        att.content_text = merged_text
        att.truncated = att.truncated or truncated
        att.metadata["has_vision"] = True
        att.metadata["vision_chars"] = len(desc)
